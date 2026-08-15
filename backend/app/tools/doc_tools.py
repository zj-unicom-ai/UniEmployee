"""解决方案文档生成工具：generate_solution_doc —— 生成标准化的 Word 方案文档。

供客户经理数字员工使用，在拜访客户后根据客户信息和推荐产品生成方案文档。
文档包含：客户信息、需求分析、推荐产品方案表格、方案优势、售后服务承诺。
"""
import os
import time
from datetime import datetime

from docx import Document                     # python-docx：Word 文档操作库
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from langchain.tools import tool
from langgraph.config import get_config

from app.paths import WORKSPACE_DATA

DATA_DIR = WORKSPACE_DATA        # 文档输出目录


@tool
def generate_solution_doc(
    customer_name: str,
    company: str = "",
    analysis: str = "",
    products: str = "",
    total_amount: str = "",
    notes: str = "",
) -> str:
    """【生成方案文档】根据客户信息和推荐产品生成标准化的 Word 解决方案文档。

    客户经理完成客户拜访后调用此工具生成可交付的方案文档。

    参数:
        customer_name: 客户姓名/称呼
        company: 客户所属企业名称（可选）
        analysis: 客户需求分析说明文字
        products: 推荐产品清单，每行格式：产品名 数量 单价 小计（空格分隔）
        total_amount: 方案总金额（字符串，如 "9788.00"）
        notes: 备注信息（可选）
    返回:
        生成的 Word 文档文件路径
    """
    # 获取当前用户 ID，用于按用户隔离文档目录
    user_id = "default"
    try:
        cfg = get_config() or {}
        user_id = (cfg.get("configurable") or {}).get("user_id", "default")
    except Exception:
        pass

    # 用户专属文档目录（workspace/data/<user_id>/）
    user_dir = DATA_DIR / user_id
    user_dir.mkdir(parents=True, exist_ok=True)

    # 生成唯一文件名：解决方案_客户名_时间戳.docx
    ts = time.strftime("%Y%m%d_%H%M%S")
    filename = f"解决方案_{customer_name}_{ts}.docx"
    filepath = user_dir / filename

    doc = Document()

    # ---- 页面设置：A4 纸，2.5cm 页边距 ----
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ---- 封面标题 ----
    title = doc.add_heading("", level=0)
    run = title.add_run(f"智选智能硬件 · 客户解决方案")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x73)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- 副标题：客户名称 + 企业 ----
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{customer_name}{'（' + company + '）' if company else ''}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x47, 0x68, 0x9A)

    # ---- 日期 ----
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # 空行

    # ---- 一、客户信息 ----
    doc.add_heading("一、客户信息", level=1)
    p = doc.add_paragraph()
    p.add_run(f"客户名称：").bold = True
    p.add_run(f"{customer_name}")
    if company:
        p2 = doc.add_paragraph()
        p2.add_run(f"所属企业：").bold = True
        p2.add_run(f"{company}")

    # ---- 二、需求分析 ----
    if analysis:
        doc.add_heading("二、需求分析", level=1)
        doc.add_paragraph(analysis)

    # ---- 三、推荐产品方案（带表格） ----
    doc.add_heading("三、推荐产品方案", level=1)
    if products:
        # 按行解析产品清单
        lines = [l.strip() for l in products.strip().split("\n") if l.strip()]
        if lines:
            # 创建表格：表头 + 数据行，4列（产品名/数量/单价/小计）
            table = doc.add_table(rows=1 + len(lines), cols=4)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表头行
            headers = ["产品名称", "数量", "单价（元）", "小计（元）"]
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            # 数据行
            for idx, line in enumerate(lines):
                parts = line.split()
                for i, part in enumerate(parts):
                    if i < 4:
                        table.rows[idx + 1].cells[i].text = part
        else:
            doc.add_paragraph(products)

    # ---- 方案总金额（醒目红色） ----
    if total_amount:
        p = doc.add_paragraph()
        p.add_run("\n方案总金额：").bold = True
        run = p.add_run(f"¥{total_amount}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # ---- 四、方案优势（固定模板） ----
    doc.add_heading("四、方案优势", level=1)
    advantages = [
        "全系产品原厂正品，享受完整质保服务",
        "支持企业批量采购优惠，量大价优",
        "免费上门安装调试（限市区范围内）",
        "7×24 小时售后技术支持",
        "提供 15 天价保服务，采购无忧",
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style="List Bullet")

    # ---- 五、售后服务承诺（固定模板） ----
    doc.add_heading("五、售后服务承诺", level=1)
    doc.add_paragraph("本方案所含产品均享受以下售后服务：")
    services = [
        "整机保修 1 年，主要部件保修 2 年",
        "保修期内非人为损坏免费维修",
        "客服热线：400-800-1234（工作日 9:00-21:00）",
        "维修寄修地址：深圳市南山区科技园南区 A3 栋 2F",
    ]
    for s in services:
        doc.add_paragraph(s, style="List Bullet")

    # ---- 页脚声明 ----
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("— 本方案由智选智能硬件客户经理自动生成，仅供参考 —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.italic = True

    # 保存文档到用户专属目录
    doc.save(str(filepath))
    return f"解决方案文档已生成：{filepath}"
